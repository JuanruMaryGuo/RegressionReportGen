from docx.oxml.ns import qn
from docx import Document

import pandas as pd

import numpy as np
from docx.shared import Inches

def generateDoc(ar, ac, dict_):
    document = Document()
    document.styles['Normal'].font.name = u'等线'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    # 插入一级标题
    document.add_heading('回归报告分析', level=0)  # 插入标题

    # 作者信息等
    p = document.add_paragraph('''
            完成人：Juanru Guo 
    ''', ) # (请在这里写上您的名字)

    # 第一部分
    document.add_heading('一、实验目的', level=1)

    # 插入段落
    document.add_paragraph('利用python进行回归分析，得到预测结果并且进行详细的统计分析。', style='List Bullet')
    document.add_paragraph('对结果进行分析并提出改进建议', style='List Bullet')

    # 第二部分
    document.add_heading('二、实验内容', level=1)

    document.add_paragraph('对数据进行相关性分析，并且进行数据可视化', style='List Bullet')
    document.add_paragraph('建立目标值与所有变量之间的多元回归方程，并观察显著性检验F和t检验。', style='List Bullet')

    # 第三部分
    document.add_heading('三、相关性分析以及数据可视化')
    document.add_paragraph('首先对数据进行相关性分析，并画出热图，颜色越深，则相关性越高')
    document.add_picture('heatmap.png', width=Inches(5))
    p = document.add_paragraph('由图中可以看出，变量')

    for i in range(0, len(ar.pvalues)):
        t = 0
        if np.abs(ac.iat[i, 0]) >= 0.8:
            if t != 0:
                p.add_run('、')
            if i == 0:
                p.add_run('常数')
                t = 1
            else:
                t = 1
                p.add_run('x')
                p.add_run(str(i))

    p.add_run('与y有很强的相关性，他们对回归的影响很可能非常显著。')
    document.add_paragraph('各个变量之间的散点图可以个更加清楚地描绘这层关系：')
    document.add_picture('set.png', width=Inches(5))

    # 第四部分
    document.add_heading('四、回归报告')
    document.add_paragraph('综合分析', style="List Bullet")
    document.add_paragraph(str(ar.summary()))
    p0 = document.add_paragraph('可得等式：y = ')
    for i in range(len(ar.pvalues)):
        if i == 0:
            p0.add_run('{0:.3f}'.format(ar.params[i]))
        else:
            p0.add_run(' + ')
            p0.add_run('{0:.3f}'.format(ar.params[i]))
            p0.add_run(' * x')
            p.add_run(str(i))
    document.add_paragraph('R方', style="List Bullet")

    pr = document.add_paragraph('R方值为：')
    pr.add_run('{0:.3f}'.format(ar.rsquared))
    pr.add_run(',调整R方值为：')
    pr.add_run('{0:.3f}'.format(ar.rsquared_adj))
    if ar.rsquared_adj >= 0.8:
        pr.add_run('，模型拟合较好。')
    elif ar.rsquared_adj >= 0.6:
        pr.add_run('，模型拟合一般，有待提升。')

    document.add_paragraph('T检验', style="List Bullet")

    p1 = document.add_paragraph('由图中可以看出，变量')
    for i in range(len(ar.pvalues)):
        t = 0
        if ar.pvalues[i] < 0.05:
            if t != 0:
                p1.add_run('、x')
                t = t + 1
            p1.add_run('x' + str(i + 1) + ', ')

    p1.add_run('通过了t检验，即认为其系数不为0，对于剩余变量，建议进行删除处理。')

    p2 = document.add_paragraph('F检验的p-value为:')

    p2.add_run('{:.4e}'.format(ar.f_pvalue))
    if ar.f_pvalue <= 0.05:
        p2.add_run('，通过了F检验，可以认为至少有一个变量的系数不为0。')
    else:
        p2.add_run('，没有通过F检验，认为所有有一个变量的系数都为0，此时模型非常不好，可能是是各个变量与y没有太强的相关性。')

    #第五部分
    if dict_ != dict():
        document.add_heading('五、机器学习模型结果')
        document.add_paragraph('各个模型平均均方误差对比结果表')
        df = pd.DataFrame(dict_)
        t = document.add_table(df.shape[0] + 1, df.shape[1])
        t.style = 'Table Grid'

        # add the header rows.
        for j in range(df.shape[-1]):
            t.cell(0, j).text = df.columns[j]

        # add the rest of the data frame
        for i in range(df.shape[0]):
            for j in range(df.shape[-1]):
                t.cell(i + 1, j).text = str(df.values[i, j])

    # 保存文档
    document.save('回归分析报告.docx')
    print('Document Generation Done.')